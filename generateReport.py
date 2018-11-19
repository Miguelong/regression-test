#!/usr/bin/python
# -*- coding: UTF-8 -*-

from  docx import  Document
from  docx.shared import  Pt
from  docx.oxml.ns import  qn
from docx.shared import Inches

document = Document()

document.add_heading('Report',0)
paragraph = document.add_paragraph('')
run = paragraph.add_run('begin: 2018-07-14 10:14:27\n')
run.font.size=Pt(12)
run = paragraph.add_run('end: 2018-07-14 10:15:09')
run.font.size=Pt(12)
document.add_heading('Testing Parameters',1)
paragraph = document.add_paragraph('')
run = paragraph.add_run('Concurrent threads: 10\n')
run.font.size=Pt(12)
run = paragraph.add_run('Url: http://10.100.34.73:8082/v1/getComparedAPI/')
run.font.size=Pt(12)
document.add_heading('Statistics',1)
document.add_picture('/Users/miguel/out.png')
paragraph = document.add_paragraph('')
run = paragraph.add_run('Get Success: 479\n')
run.font.size=Pt(12)
run = paragraph.add_run('max:2.14 min:0.48 average:0.510\n\n')
run.font.size=Pt(12)
run = paragraph.add_run('[ 0.48 , 0.98 ): 475\n')
run.font.size=Pt(12)
run = paragraph.add_run('(0.48,39),(0.49,270),(0.5,112),(0.51,15),(0.52,2),(0.53,4),(0.54,3),(0.55,3),(0.56,6),(0.57,3),(0.58,2),(0.59,1),(0.6,3),(0.61,8),(0.63,1),(0.64,1),(0.65,1),(0.88,1)\n')
run.font.size=Pt(12)

document.add_heading('Fail ',1)
paragraph = document.add_paragraph('')
run = paragraph.add_run('total:23')
run.font.size=Pt(12)
document.add_heading('Code 500',2)
paragraph = document.add_paragraph('')
run = paragraph.add_run('total:21\n')
run.font.size=Pt(12)
run = paragraph.add_run('0000043237\n')
run.font.size=Pt(12)

document.add_heading("ValueError('No JSON object could be decoded',)",2)
paragraph = document.add_paragraph('')
run = paragraph.add_run('total:2\n')
run.font.size=Pt(12)
run = paragraph.add_run('0000043237\n')
run.font.size=Pt(12)

document.save('demo.docx')
